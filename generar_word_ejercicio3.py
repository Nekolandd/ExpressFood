# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\diani\Documents\ExpressFood\Examen_Diseno_Ejercicio3.docx"


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, fill="D9E2F3"):
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    style_header_row(table.rows[0])

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row[c_idx].text = str(val)
            for p in row[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


doc = Document()

# Título
title = doc.add_heading("Plan de Trabajo — Clínica SaludPlus", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    "Proyecto: Aplicación móvil para gestión de citas médicas, consulta de "
    "información clínica y notificaciones."
)
p.runs[0].font.size = Pt(11)

doc.add_heading("Fases, tiempos y entregables", level=2)

fases_headers = ["Fase", "Tiempo", "Actividad", "Entregables"]
fases_rows = [
    [
        "1. Análisis",
        "1 sem",
        "Levantamiento de requisitos, perfiles de pacientes y definición de alcance",
        "Documento de requisitos, mapa de pantallas (9 módulos), casos de uso",
    ],
    [
        "2. Diseño UX/UI",
        "2 sem",
        "Esquema, flujo de navegación, guía de estilo (colores, tipografía, componentes)",
        "Esquema de baja/media fidelidad, guía visual, arquitectura de información",
    ],
    [
        "3. Prototipado",
        "1 sem",
        "Prototipo interactivo de las 9 pantallas y validación con el cliente",
        "Prototipo en Figma, ajustes según feedback",
    ],
    [
        "4. Accesibilidad e internacionalización",
        "1 sem",
        "Adaptar la app para personas con discapacidad (contraste, botones grandes, "
        "lectores de pantalla) y prepararla para varios idiomas y países "
        "(español/inglés, fechas, horas, monedas)",
        "Checklist de accesibilidad, archivos de traducción, guía de formatos regionales",
    ],
    [
        "5. Desarrollo",
        "3 sem",
        "Implementación móvil (Android/iOS), login, citas, resultados, directorio, "
        "perfil y notificaciones",
        "App funcional con las 9 pantallas requeridas",
    ],
    [
        "6. Pruebas",
        "1 sem",
        "Usabilidad, accesibilidad, soporte multidioma, seguridad básica "
        "(login, datos clínicos)",
        "Informe de pruebas, correcciones aplicadas",
    ],
    [
        "7. Entrega y cierre",
        "1 sem",
        "Documentación, capacitación y despliegue",
        "Manual de usuario, manual técnico, app publicada o APK/IPA de prueba",
    ],
]

add_table(doc, fases_headers, fases_rows, [1.1, 0.7, 2.4, 2.0])

doc.add_paragraph()
bold = doc.add_paragraph()
run = bold.add_run("Duración total estimada: 10 semanas")
run.bold = True
run.font.size = Pt(11)

doc.add_page_break()
doc.add_heading("Cronograma general", level=2)
doc.add_heading("Vista por semanas", level=3)

cron_headers = ["Semana", "Fase", "Actividades principales", "Hito"]
cron_rows = [
    ["1", "Análisis", "Reuniones con cliente, casos de uso, mapa de 9 pantallas", "Requisitos aprobados"],
    ["2", "Diseño UX/UI", "Wireframes, flujo de navegación, arquitectura de información", "Wireframes v1"],
    ["3", "Diseño UX/UI", "Guía de estilo, componentes (botones, formularios, tarjetas)", "Guía visual aprobada"],
    ["4", "Prototipado", "Prototipo interactivo + validación con cliente", "Prototipo validado"],
    [
        "5",
        "Accesibilidad e internacionalización",
        "Contraste y tipografía legible; etiquetas en formularios; soporte para lectores "
        "de pantalla; textos en español e inglés; formatos de fecha, hora y moneda",
        "Guía de accesibilidad e internacionalización",
    ],
    ["6", "Desarrollo", "Login, pantalla principal, agenda y detalle de cita", "Módulo de citas"],
    ["7", "Desarrollo", "Mis citas, resultados médicos, directorio médico", "Módulo clínico"],
    ["8", "Desarrollo", "Perfil, notificaciones, ajustes finales Android/iOS", "App funcional"],
    ["9", "Pruebas", "Usabilidad, accesibilidad, idiomas y formatos regionales, seguridad", "Informe de pruebas"],
    ["10", "Entrega", "Documentación, capacitación, despliegue", "Proyecto entregado"],
]

add_table(doc, cron_headers, cron_rows, [0.6, 1.4, 2.8, 1.5])

doc.add_paragraph()
doc.add_heading("Diagrama por semanas", level=3)

leg = doc.add_paragraph("Leyenda: ■ = actividad en curso en esa semana")
leg.runs[0].font.size = Pt(10)
leg2 = doc.add_paragraph("S#: Semana en la que se encuentra")
leg2.runs[0].font.size = Pt(10)

gantt_headers = ["Fase", "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "S9", "S10"]
gantt_rows = [
    ["1. Análisis", "■", "", "", "", "", "", "", "", "", ""],
    ["2. Diseño UX/UI", "", "■", "■", "", "", "", "", "", "", ""],
    ["3. Prototipado", "", "", "", "■", "", "", "", "", "", ""],
    ["4. Accesibilidad e internacionalización", "", "", "", "", "■", "", "", "", "", ""],
    ["5. Desarrollo", "", "", "", "", "", "■", "■", "■", "", ""],
    ["6. Pruebas", "", "", "", "", "", "", "", "", "■", ""],
    ["7. Entrega y cierre", "", "", "", "", "", "", "", "", "", "■"],
]

gantt_table = add_table(doc, gantt_headers, gantt_rows, [1.8] + [0.35] * 10)

for r_idx in range(1, len(gantt_table.rows)):
    for c_idx in range(1, 11):
        cell = gantt_table.rows[r_idx].cells[c_idx]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.text.strip() == "■":
            set_cell_shading(cell, "4472C4")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()
dep = doc.add_paragraph(
    "Cada fase depende de la anterior: no se desarrolla sin prototipo aprobado; "
    "no se entrega sin pruebas de usabilidad, accesibilidad e internacionalización."
)
dep.runs[0].font.size = Pt(11)

doc.add_heading("Justificación de la secuencia", level=2)

just = doc.add_paragraph(
    "La secuencia del plan sigue el proceso de Ingeniería de la Interfaz: primero se "
    "analizan requisitos y usuarios (pacientes con poca experiencia tecnológica), luego "
    "se diseña y valida con el cliente una interfaz intuitiva, profesional y de "
    "navegación simple antes de programar, lo que reduce retrabajo. La fase de "
    "accesibilidad e internacionalización se define antes del desarrollo para cumplir "
    "contraste, tipografías legibles, botones adecuados, lectores de pantalla, etiquetas "
    "en formularios, soporte en español e inglés, formatos de fecha/hora y monedas. "
    "El desarrollo implementa las 9 pantallas requeridas en Android e iOS; las pruebas "
    "verifican usabilidad, accesibilidad, multidioma y seguridad de los datos clínicos; "
    "y la entrega cierra con documentación y despliegue. Este orden permite cumplir el "
    "objetivo del cliente — gestionar citas, consultar información clínica y recibir "
    "notificaciones de forma sencilla y segura — con entregables claros en cada fase y "
    "un cronograma de 10 semanas controlable mediante hitos semanales."
)
just.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
just.runs[0].font.size = Pt(11)

doc.save(OUTPUT)
print(f"Documento creado: {OUTPUT}")
